"""The ITIL baseline: every practice's deliverables, each compiled from the generators and run on the real inputs.

The register lives in proofs/itil.yaml as data.  A deliverable is a chain of operations; every
operation is bound here to the library functions that perform it.  Nothing in this file decides
what a deliverable is: it executes the register and records what came out.
"""
from __future__ import annotations

import collections
import hashlib
import json
import re
from pathlib import Path
from typing import Any

import yaml

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "proofs" / "out" / "itil"

Rows = list[dict[str, Any]]
IDENT = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")

# every operation, and the functions it is bound to; printed on the arrow lines
BINDINGS: dict[str, str] = {
    "facts": "P1: compiled_ai read, parse, extract, project, route, check, adjudicate",
    "required_actions": "P4: the obligatory facts with agent quotes",
    "ordered_steps": "P2: z3 consistency proof; networkx forced order",
    "forced_edges": "P2: networkx.transitive_reduction",
    "anchors": "P3: timexy",
    "definitions": "P1 facts where the event lemma is mean, include or define",
    "prohibitions": "P1 facts where obligatory and negated",
    "who_must_what": "P44: duckdb count per (agent, action)",
    "records": "P6: csv.DictReader; pydantic.create_model",
    "cases": "duckdb: one row per case from the event rows",
    "events": "P6 rows as a relation",
    "minutes_sentences": "P7: compiled_ai parse of the minutes; one row per sentence with its event lemmas",
    "minutes_sections": "markdown.markdown; lxml.html: heading sections and their list items",
    "participants": "markdown.markdown; lxml.html: the Present list of each meeting",
    "attendance": "P13: voting members present per meeting",
    "decisions": "P13: zen.ZenDecision.evaluate",
    "tags": "P10/P26: clingo shared-word tagging",
    "conformance": "P9: pm4py.conformance_diagnostics_token_based_replay",
    "variants": "P36: pm4py.get_variants_as_tuples",
    "dfg": "P23: pm4py.discover_dfg",
    "handover": "pm4py.discover_handover_of_work_network",
    "execution_trace": "P29: SpiffWorkflow execution of the process",
    "roundtrips": "P19, P30, P40, P50: reverse proofs and csv_diff.compare",
    "bottlenecks": "P46: duckdb join of edges and waiting time",
    "filter": "duckdb WHERE",
    "select": "duckdb SELECT",
    "derive": "duckdb expression (closed set: month, days_between, hours_between, contains, ratio, length)",
    "group": "duckdb GROUP BY with count, sum, avg, min, max, list",
    "join": "duckdb JOIN",
    "sort": "duckdb ORDER BY",
    "limit": "duckdb LIMIT",
    "distinct": "duckdb DISTINCT",
    "pivot": "duckdb GROUP BY, laid out as a matrix",
    "save_as": "the relation kept under a name for a later join",
    "clingo": "clingo.Control.add / ground / solve; rows as ground facts",
    "decision_table": "jinja2 -> GoRules JDM; zen.ZenEngine.create_decision; zen.ZenDecision.evaluate",
    "workbook": "openpyxl.Workbook.save",
    "document": "docx.Document; add_table; docx.document.Document.save",
    "checklist": "docx.Document; one box glyph per row; docx.document.Document.save",
    "page": "plotly.express / plotly.graph_objects; plotly.io.write_html | jinja2 table page",
    "bpmn": "pm4py.objects.bpmn.obj.BPMN with parallel gateways; bpmn exporter; canonical ids",
    "json": "json.dumps sort_keys",
}


OFFICE = (".xlsx", ".docx", ".pptx")
FIXED_STAMP = "2000-01-01T00:00:00Z"


def canonical_office(path: Path) -> None:
    """Make an Office file byte-reproducible: fixed core timestamps, fixed zip entry dates, sorted entries."""
    import zipfile

    if path.suffix.lower() not in OFFICE:
        return
    with zipfile.ZipFile(path) as z:
        entries = {i.filename: z.read(i.filename) for i in z.infolist()}
    core = entries.get("docProps/core.xml")
    if core is not None:
        text = core.decode("utf-8")
        text = re.sub(r"(<dcterms:created[^>]*>)[^<]*(</dcterms:created>)", rf"\g<1>{FIXED_STAMP}\g<2>", text)
        text = re.sub(r"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)", rf"\g<1>{FIXED_STAMP}\g<2>", text)
        entries["docProps/core.xml"] = text.encode("utf-8")
    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name in sorted(entries):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            z.writestr(info, entries[name])
    tmp.replace(path)


# ---- the register of proven digests ----------------------------------------------------------
# A deliverable is proven once.  Its digest is registered, and every later run compares the new
# digest against the registered one instead of proving it again.
REGISTER_PATH = ROOT / "proofs" / "register.json"
REGISTER: dict[str, str] = {}
NEW: dict[str, str] = {}
CHANGED: dict[str, tuple[str, str]] = {}


def load_register() -> None:
    REGISTER.clear()
    NEW.clear()
    CHANGED.clear()
    if REGISTER_PATH.exists():
        REGISTER.update(json.loads(REGISTER_PATH.read_text(encoding="utf-8")))


def check_register(rel: str, sha: str) -> str:
    if rel in REGISTER:
        if REGISTER[rel] == sha:
            return "reproduced"
        CHANGED[rel] = (REGISTER[rel], sha)
        return "CHANGED"
    NEW[rel] = sha
    return "registered"


def save_register(reregister: bool = False) -> None:
    REGISTER.update(NEW)
    if reregister:
        REGISTER.update({k: v[1] for k, v in CHANGED.items()})
    REGISTER_PATH.write_text(json.dumps(dict(sorted(REGISTER.items())), indent=1) + "\n", encoding="utf-8")


def sha256_file(p: Path) -> str:
    return hashlib.sha256(p.read_bytes()).hexdigest()


def slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")


def ident(name: str) -> str:
    n = re.sub(r"[^A-Za-z0-9_]+", "_", name).strip("_").lower()
    if not IDENT.match(n):
        n = "f_" + n
    return n


class Ctx:
    """What the operations can reach: the results of the numbered proofs, plus lazily built relations."""

    def __init__(self, results: dict[Any, Any], helpers: dict[str, Any]) -> None:
        self.results = results
        self.h = helpers
        self.cache: dict[str, Rows] = {}

    # ---- sources -------------------------------------------------------------------------

    def facts(self, pack: str) -> Rows:
        c = self.results[("facts", pack)]
        sent = {ps.sentence.id: ps.sentence.text for ps in c["parsed"]}
        return [
            {"id": a.id, "predicate": a.predicate, "arg1": a.args[0], "arg2": a.args[1] if len(a.args) > 1 else "", "sentence_id": a.sentence_id, "quote": a.quote, "sentence": sent.get(a.sentence_id, ""), "unit": a.sentence_id.split(":")[1]}
            for a in c["atoms"]
        ]

    def required_actions(self, pack: str) -> Rows:
        c = self.results[("policy", pack)]
        sent = {ps.sentence.id: ps.sentence.text for ps in c["parsed"]}
        agent = collections.defaultdict(list)
        for a in c["atoms"]:
            if a.predicate == "agent":
                agent[a.args[0]].append(a.quote)
        return [{"event": r["event"], "action": r["action"], "who": "; ".join(agent.get(r["event"], [])), "sentence_id": r["sentence_id"], "sentence": sent.get(r["sentence_id"], ""), "quote": r["quote"], "unit": r["sentence_id"].split(":")[1]} for r in c["required"]]

    def ordered_steps(self, pack: str) -> Rows:
        import networkx as nx

        c = self.results[("ordered steps", pack)]
        o = c["ordering"]
        lemma = {a.args[0]: a.args[1] for a in c["atoms"] if a.predicate == "event"}
        sent = {ps.sentence.id: ps.sentence.text for ps in c["parsed"]}
        G = nx.DiGraph()
        G.add_nodes_from(o.order)
        G.add_edges_from(o.forced)
        layer = {n: i for i, gen in enumerate(nx.topological_generations(G)) for n in gen}
        return [{"position": i, "event": e, "step": lemma.get(e, "?"), "phase": layer.get(e, 0) + 1, "sentence": sent.get(e.split("#")[0], ""), "unit": e.split(":")[1]} for i, e in enumerate(o.order, 1)]

    def forced_edges(self, pack: str) -> Rows:
        c = self.results[("ordered steps", pack)]
        lemma = {a.args[0]: a.args[1] for a in c["atoms"] if a.predicate == "event"}
        return [{"before": a, "after": b, "before_step": lemma.get(a, "?"), "after_step": lemma.get(b, "?")} for a, b in c["ordering"].forced]

    def anchors(self, pack: str) -> Rows:
        return [dict(r) for r in self.results[("anchor dates", pack)]["anchors"]]

    def definitions(self, pack: str) -> Rows:
        rows = self.facts(pack)
        defs = {r["arg1"] for r in rows if r["predicate"] == "event" and r["arg2"] in ("mean", "include", "define")}
        seen: set[str] = set()
        out = []
        for r in rows:
            if r["predicate"] == "event" and r["arg1"] in defs and r["sentence_id"] not in seen:
                seen.add(r["sentence_id"])
                out.append({"event": r["arg1"], "verb": r["arg2"], "sentence_id": r["sentence_id"], "sentence": r["sentence"], "unit": r["unit"]})
        return out

    def prohibitions(self, pack: str) -> Rows:
        rows = self.facts(pack)
        obl = {r["arg1"] for r in rows if r["predicate"] == "obligatory"}
        neg = {r["arg1"] for r in rows if r["predicate"] == "negated"}
        return [{"event": r["arg1"], "action": r["arg2"], "sentence_id": r["sentence_id"], "sentence": r["sentence"], "unit": r["unit"]} for r in rows if r["predicate"] == "event" and r["arg1"] in obl and r["arg1"] in neg]

    def who_must_what(self, pack: str) -> Rows:
        rows = self.required_actions(pack)
        c = collections.Counter((r["who"].lower(), r["action"]) for r in rows if r["who"])
        return [{"who": w, "action": a, "count": n} for (w, a), n in sorted(c.items(), key=lambda kv: (-kv[1], kv[0]))]

    def records(self, name: str) -> Rows:
        return [{ident(k): v for k, v in r.items()} for r in self.results[("records", name)]["rows"]]

    def events(self, name: str = "receipt-csv") -> Rows:
        if "events" not in self.cache:
            rows = self.results[("records", name)]["rows"]
            self.cache["events"] = [
                {"case": r["case:concept:name"], "activity": r["concept:name"], "ts": r["time:timestamp"], "resource": r.get("org:resource", ""), "group": r.get("org:group", ""), "department": r["case:department"], "channel": r["case:channel"], "deadline": r["case:deadline"], "enddate": r["case:enddate"], "startdate": r.get("case:startdate", ""), "responsible": r.get("case:responsible", "")}
                for r in rows
            ]
        return list(self.cache["events"])

    def cases(self, name: str = "receipt-csv") -> Rows:
        if "cases" not in self.cache:
            import duckdb

            import pandas as pd

            con = duckdb.connect()
            df = pd.DataFrame([(e["case"], e["activity"], e["ts"], e["resource"], e["department"], e["channel"], e["deadline"], e["enddate"], e["responsible"]) for e in self.events(name)], columns=["case_id", "activity", "ts_text", "resource", "department", "channel", "deadline", "enddate", "responsible"])
            con.register("ev_raw", df)
            con.execute("create table ev as select case_id, activity, cast(ts_text as timestamptz) as ts, resource, department, channel, deadline, enddate, responsible from ev_raw")
            end = con.execute("select max(ts) from ev").fetchone()[0]
            res = con.execute(
                """
                select case_id, any_value(department), any_value(channel), any_value(responsible), min(ts), max(ts), count(*),
                       round(epoch(max(ts) - min(ts))/86400, 2), any_value(deadline), any_value(enddate),
                       case when any_value(enddate) = '' then 'open'
                            when try_cast(any_value(enddate) as timestamptz) > try_cast(any_value(deadline) as timestamptz) then 'after deadline'
                            else 'by deadline' end,
                       strftime(min(ts), '%Y-%m'), round(epoch(? - max(ts))/86400, 2)
                from ev group by case_id order by case_id
                """,
                [end],
            ).fetchall()
            cols = ["case", "department", "channel", "responsible", "start", "end", "events", "duration_days", "deadline", "enddate", "outcome", "month", "days_since_last_event"]
            self.cache["cases"] = [dict(zip(cols, r)) for r in res]
        return list(self.cache["cases"])

    def minutes_sentences(self, _: str = "") -> Rows:
        c = self.results[("parsed minutes", "nodejs-tsc-minutes")]
        lem = collections.defaultdict(list)
        for a in c["atoms"]:
            if a.predicate == "event":
                lem[a.sentence_id].append(a.args[1])
        return [{"meeting": ps.sentence.id.split(":")[0], "sentence_id": ps.sentence.id, "sentence": ps.sentence.text, "lemmas": " ".join(sorted(set(lem.get(ps.sentence.id, []))))} for ps in c["parsed"]]

    def minutes_sections(self, _: str = "") -> Rows:
        if "sections" not in self.cache:
            import lxml.html
            import markdown

            out = []
            for src in self.results[("parsed minutes", "nodejs-tsc-minutes")]["sources"]:
                html = markdown.markdown(Path(src.path).read_text(encoding="utf-8"))
                doc = lxml.html.fromstring(html)
                section = ""
                for el in doc.iter():
                    if not isinstance(el.tag, str):
                        continue
                    if el.tag in ("h1", "h2", "h3", "h4"):
                        section = " ".join(el.text_content().split())
                    elif el.tag == "li" and el.getparent() is not None and el.getparent().tag == "ul" and not any(p.tag == "li" for p in el.iterancestors()):
                        text = " ".join(el.text_content().split())
                        links = [a.get("href") for a in el.iter("a") if a.get("href")]
                        out.append({"meeting": src.id, "section": section, "item": text[:300], "link": links[0] if links else "", "links": len(links)})
            self.cache["sections"] = out
        return list(self.cache["sections"])

    def participants(self, _: str = "") -> Rows:
        if "participants" not in self.cache:
            import lxml.html
            import markdown

            out = []
            for src in self.results[("parsed minutes", "nodejs-tsc-minutes")]["sources"]:
                html = markdown.markdown(Path(src.path).read_text(encoding="utf-8"))
                doc = lxml.html.fromstring(html)
                for h in doc.iter("h2"):
                    if h.text_content().strip() == "Present":
                        ul = next(s for s in h.itersiblings() if s.tag == "ul")
                        for li in ul.iter("li"):
                            t = " ".join(li.text_content().split())
                            m = re.match(r"^(.*?)\s*(?:@(\S+))?\s*\((.*?)\)\s*$", t)
                            name, handle, role = (m.group(1), m.group(2) or "", m.group(3)) if m else (t, "", "")
                            out.append({"meeting": src.id, "name": name.strip(), "handle": handle, "role": role})
            self.cache["participants"] = out
        return list(self.cache["participants"])

    def attendance(self, _: str = "") -> Rows:
        return [dict(r) for r in self.results[("decisions", "majority")]["attendance"]]

    def decisions(self, _: str = "majority") -> Rows:
        rows = self.results[("decisions", "majority")]["decision_rows"]
        return [dict(zip(rows[0], r)) for r in rows[1:]]

    def tags(self, pack: str) -> Rows:
        c = self.results[("tagged steps", pack)]
        sents = {r["sentence_id"]: r["sentence"] for r in self.minutes_sentences()}
        return [{"meeting": t["sentence_id"].split(":")[0], "sentence_id": t["sentence_id"], "sentence": sents.get(t["sentence_id"], ""), "step": t["step"], "step_lemma": c["lemma"].get(t["step"], "?"), "shared": " ".join(t["shared"]), "tie": t["tie"]} for t in c["tags"]]

    def conformance(self, log: str = "receipt-xes") -> Rows:
        rows = self.results[("conformance", log)]["conformance_rows"]
        return [dict(zip(["case", "fitness", "is_fit", "missing_tokens", "remaining_tokens"], r)) for r in rows[1:]]

    def variants(self, log: str = "receipt-xes") -> Rows:
        import pm4py

        if "variants" not in self.cache:
            v = pm4py.get_variants_as_tuples(self.results[("log", log)]["log"])
            items = sorted(((k, (n if isinstance(n, int) else len(n))) for k, n in v.items()), key=lambda kv: (-kv[1], kv[0]))
            self.cache["variants"] = [{"rank": i, "cases": n, "length": len(k), "variant": " -> ".join(k), "first": k[0], "last": k[-1]} for i, (k, n) in enumerate(items, 1)]
        return list(self.cache["variants"])

    def dfg(self, log: str = "receipt-xes") -> Rows:
        rows = self.results[("dfg", log)]["dfg_rows"]
        return [{"from": a, "to": b, "count": n} for a, b, n in rows[1:]]

    def handover(self, log: str = "receipt-xes") -> Rows:
        import pm4py

        if "handover" not in self.cache:
            net = pm4py.discover_handover_of_work_network(self.results[("log", log)]["log"])
            conn = net.connections if hasattr(net, "connections") else net
            self.cache["handover"] = [{"from": a, "to": b, "value": round(float(v), 6)} for (a, b), v in sorted(conn.items(), key=lambda kv: (-kv[1], kv[0]))]
        return list(self.cache["handover"])

    def execution_trace(self, pack: str) -> Rows:
        return [{"position": i, "task": t} for i, t in enumerate(self.h["traces"][pack], 1)]

    def roundtrips(self, _: str = "") -> Rows:
        return [dict(r) for r in self.h["roundtrips"]]

    def bottlenecks(self, _: str = "") -> Rows:
        return [dict(r) for r in self.h["bottlenecks"]]

    def source(self, op: str, arg: str) -> Rows:
        return getattr(self, op)(arg) if arg else getattr(self, op)()


# ---- relational operations on rows, through duckdb ------------------------------------------------


def _con(rows: Rows, name: str = "r") -> Any:
    import duckdb

    con = duckdb.connect()
    if not rows:
        con.execute(f"create table {name}(empty varchar)")
        return con
    cols = list(rows[0].keys())
    for c in cols:
        assert IDENT.match(c), c
    import pandas as pd

    con.register(name, pd.DataFrame(rows, columns=cols))
    return con


def _fetch(con: Any, sql: str, params: list[Any] | None = None) -> Rows:
    cur = con.execute(sql, params or [])
    cols = [d[0] for d in cur.description]
    return [dict(zip(cols, r)) for r in cur.fetchall()]


OPS = {"eq": "=", "ne": "<>", "gt": ">", "lt": "<", "ge": ">=", "le": "<="}


def op_filter(rows: Rows, where: list[list[Any]]) -> Rows:
    if not rows:
        return rows
    con = _con(rows)
    clauses, params = [], []
    for field, op, value in where:
        assert IDENT.match(field)
        if op in OPS:
            clauses.append(f'"{field}" {OPS[op]} ?')
            params.append(value)
        elif op == "contains":
            clauses.append(f'lower(cast("{field}" as varchar)) like ?')
            params.append(f"%{str(value).lower()}%")
        elif op == "in":
            clauses.append(f'"{field}" in ({",".join("?" for _ in value)})')
            params.extend(value)
        elif op == "isnull":
            clauses.append(f'("{field}" is null or cast("{field}" as varchar) = \'\')')
        elif op == "notnull":
            clauses.append(f'("{field}" is not null and cast("{field}" as varchar) <> \'\')')
        elif op == "matches":
            clauses.append(f'regexp_matches(cast("{field}" as varchar), ?)')
            params.append(value)
        else:
            raise ValueError(op)
    return _fetch(con, f"select * from r where {' and '.join(clauses)}", params)


def op_select(rows: Rows, fields: list[str]) -> Rows:
    return [{f: r.get(f) for f in fields} for r in rows]


def op_derive(rows: Rows, field: str, fn: str, **kw: Any) -> Rows:
    if not rows:
        return rows
    con = _con(rows)
    for k, v in kw.items():
        if k in ("of", "a", "b", "num", "den"):
            assert IDENT.match(v), v
    expr = {
        "month": lambda: f"strftime(try_cast(\"{kw['of']}\" as timestamptz), '%Y-%m')",
        "days_between": lambda: f"round(epoch(try_cast(\"{kw['b']}\" as timestamptz) - try_cast(\"{kw['a']}\" as timestamptz))/86400, 2)",
        "hours_between": lambda: f"round(epoch(try_cast(\"{kw['b']}\" as timestamptz) - try_cast(\"{kw['a']}\" as timestamptz))/3600, 2)",
        "contains": lambda: f"contains(lower(cast(\"{kw['of']}\" as varchar)), '{str(kw['value']).lower().replace(chr(39), '')}')",
        "ratio": lambda: f"round(cast(\"{kw['num']}\" as double) / nullif(cast(\"{kw['den']}\" as double), 0), 4)",
        "length": lambda: f"length(cast(\"{kw['of']}\" as varchar))",
        "words": lambda: f"length(cast(\"{kw['of']}\" as varchar)) - length(replace(cast(\"{kw['of']}\" as varchar), ' ', '')) + 1",
        "upper": lambda: f"upper(cast(\"{kw['of']}\" as varchar))",
        "gt": lambda: f"(cast(\"{kw['of']}\" as double) > {float(kw['value'])})",
        "year": lambda: f"strftime(try_cast(\"{kw['of']}\" as timestamptz), '%Y')",
    }[fn]()
    assert IDENT.match(field)
    return _fetch(con, f'select *, {expr} as "{field}" from r')


AGG = {
    "count": "count(*)",
    "sum": "sum(cast({f} as double))",
    "avg": "round(avg(cast({f} as double)), 2)",
    "min": "min({f})",
    "max": "max({f})",
    "list": "string_agg(cast({f} as varchar), ', ' order by {f})",
    "distinct": "count(distinct {f})",
    "median": "median(cast({f} as double))",
}


def op_group(rows: Rows, by: list[str], aggregates: list[list[str]]) -> Rows:
    if not rows:
        return rows
    con = _con(rows)
    for b in by:
        assert IDENT.match(b)
    parts = []
    for fn, field, alias in aggregates:
        assert IDENT.match(alias) and (field == "*" or IDENT.match(field))
        quoted = '"' + field + '"' if field != "*" else "*"
        parts.append(AGG[fn].format(f=quoted) + ' as "' + alias + '"')
    quoted_by = ['"' + b + '"' for b in by]
    sel = ", ".join(quoted_by + parts)
    group = ("group by " + ", ".join(quoted_by)) if by else ""
    order = ", ".join(quoted_by) if by else "1"
    return _fetch(con, f"select {sel} from r {group} order by {order}")


def op_join(rows: Rows, right: Rows, on: list[list[str]], how: str = "inner") -> Rows:
    if not rows or not right:
        return rows if how == "left" else []
    con = _con(rows, "l")
    import pandas as pd

    con.register("rr", pd.DataFrame(right))
    lcols = list(rows[0].keys())
    rcols = [c for c in right[0].keys() if c not in lcols]
    for a, b in on:
        assert IDENT.match(a) and IDENT.match(b)
    cond = " and ".join(f'l."{a}" = rr."{b}"' for a, b in on)
    sel = ", ".join([f'l."{c}"' for c in lcols] + [f'rr."{c}"' for c in rcols])
    return _fetch(con, f"select {sel} from l {'left' if how == 'left' else 'inner'} join rr on {cond}")


def op_sort(rows: Rows, by: list[str], desc: bool = False) -> Rows:
    for b in by:
        assert IDENT.match(b)
    return sorted(rows, key=lambda r: tuple((r.get(b) is None, r.get(b)) for b in by), reverse=desc)


def op_limit(rows: Rows, n: int) -> Rows:
    return rows[:n]


def op_distinct(rows: Rows, fields: list[str] | None = None) -> Rows:
    seen: set[Any] = set()
    out = []
    for r in rows:
        key = tuple(r.get(f) for f in (fields or list(r.keys())))
        if key not in seen:
            seen.add(key)
            out.append({f: r.get(f) for f in fields} if fields else r)
    return out


def op_pivot(rows: Rows, row: str, col: str, value: str = "*", fn: str = "count") -> Rows:
    g = op_group(rows, [row, col], [[fn, value, "v"]])
    cols = sorted({r[col] for r in g}, key=str)
    cell = {(r[row], r[col]): r["v"] for r in g}
    return [{row: rv, **{str(c): cell.get((rv, c), 0) for c in cols}} for rv in sorted({r[row] for r in g}, key=str)]


def op_clingo(rows: Rows, program: str, in_pred: str, fields: list[str], out_pred: str, out_fields: list[str]) -> Rows:
    import clingo

    def lit(v: Any) -> str:
        if isinstance(v, bool):
            return "true" if v else "false"
        if isinstance(v, int):
            return str(v)
        return '"' + str(v).replace("\\", "\\\\").replace('"', '\\"') + '"'

    prog = [f"{in_pred}({','.join(lit(r.get(f)) for f in fields)})." for r in rows] + [program, f"#show {out_pred}/{len(out_fields)}."]
    ctl = clingo.Control(["0"])
    ctl.add("base", [], "\n".join(prog))
    ctl.ground([("base", [])])
    syms: list[Any] = []
    ctl.solve(on_model=lambda m: syms.extend(m.symbols(shown=True)))
    out = []
    for s in syms:
        if s.name == out_pred:
            vals = [a.string if a.type == clingo.SymbolType.String else (a.number if a.type == clingo.SymbolType.Number else str(a)) for a in s.arguments]
            out.append(dict(zip(out_fields, vals)))
    return sorted(out, key=lambda r: tuple(str(v) for v in r.values()))


JDM = """{"contentType": "application/vnd.gorules.decision", "nodes": [
 {"id": "in", "type": "inputNode", "name": "in", "position": {"x": 0, "y": 0}},
 {"id": "dt", "type": "decisionTableNode", "name": {{ name|tojson }}, "position": {"x": 0, "y": 0}, "content": {"hitPolicy": "first",
  "inputs": [{% for i in inputs %}{"id": "i{{ loop.index }}", "name": {{ i|tojson }}, "field": {{ i|tojson }}}{{ "," if not loop.last }}{% endfor %}],
  "outputs": [{% for o in outputs %}{"id": "o{{ loop.index }}", "name": {{ o|tojson }}, "field": {{ o|tojson }}}{{ "," if not loop.last }}{% endfor %}],
  "rules": [{% for r in rules %}{"_id": "r{{ loop.index }}", "_description": {{ r.because|tojson }}{% for i in inputs %}, "i{{ loop.index }}": {{ (r.when.get(i, ""))|tojson }}{% endfor %}{% for o in outputs %}, "o{{ loop.index }}": {{ (r.then.get(o, ""))|tojson }}{% endfor %}}{{ "," if not loop.last }}{% endfor %}]}},
 {"id": "out", "type": "outputNode", "name": "out", "position": {"x": 0, "y": 0}}],
 "edges": [{"id": "e1", "sourceId": "in", "targetId": "dt", "type": "edge"}, {"id": "e2", "sourceId": "dt", "targetId": "out", "type": "edge"}]}
"""


def op_decision_table(rows: Rows, name: str, inputs: list[str], outputs: list[str], rules: list[dict[str, Any]], out_dir: Path) -> tuple[Rows, Path]:
    import jinja2
    import zen

    jdm = jinja2.Environment().from_string(JDM).render(name=name, inputs=inputs, outputs=outputs, rules=rules)
    json.loads(jdm)
    p = out_dir / "policy.jdm.json"
    p.write_text(jdm, encoding="utf-8")
    dec = zen.ZenEngine().create_decision(jdm)
    out = []
    for r in rows:
        res = dec.evaluate({i: r.get(i) for i in inputs})["result"]
        out.append({**r, **{o: res.get(o) for o in outputs}})
    return out, p


# ---- renders --------------------------------------------------------------------------------


def render_workbook(rows: Rows, path: Path, sheet: str = "rows") -> Path:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet[:31]
    if rows:
        ws.append(list(rows[0].keys()))
        for r in rows:
            ws.append([v if isinstance(v, (int, float, str)) or v is None else str(v) for v in r.values()])
    wb.save(path)
    return path


def render_document(rows: Rows, path: Path, title: str, columns: list[str] | None = None) -> Path:
    import docx

    d = docx.Document()
    d.add_heading(title, 1)
    cols = columns or (list(rows[0].keys()) if rows else [])
    t = d.add_table(rows=1, cols=len(cols))
    for i, c in enumerate(cols):
        t.rows[0].cells[i].text = c
    for r in rows:
        cells = t.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(r.get(c, ""))
    d.save(path)
    return path


def render_checklist(rows: Rows, path: Path, title: str, item: str, note: str | None = None) -> Path:
    import docx

    d = docx.Document()
    d.add_heading(title, 1)
    for r in rows:
        d.add_paragraph(f"☐  {r.get(item, '')}")
        if note and r.get(note):
            d.add_paragraph(str(r.get(note)), style="Intense Quote")
    d.save(path)
    return path


TABLE_PAGE = """<!doctype html><html><head><meta charset="utf-8"><title>{{ title }}</title>
<style>body{font-family:system-ui,sans-serif;max-width:70rem;margin:2rem auto;padding:0 1rem;color:#222}table{border-collapse:collapse;width:100%}
td,th{border:1px solid #ccc;padding:.35rem .5rem;text-align:left;vertical-align:top}th{background:#f4f4f4}</style></head><body>
<h1>{{ title }}</h1><table><tr>{% for c in cols %}<th>{{ c }}</th>{% endfor %}</tr>
{% for r in rows %}<tr>{% for c in cols %}<td>{{ r[c] }}</td>{% endfor %}</tr>
{% endfor %}</table></body></html>
"""


def render_page(rows: Rows, path: Path, title: str, kind: str = "table", x: str = "", y: str = "", color: str = "", size: str = "") -> Path:
    if kind == "table" or not rows:
        import jinja2

        cols = list(rows[0].keys()) if rows else []
        path.write_text(jinja2.Environment(autoescape=True).from_string(TABLE_PAGE).render(title=title, cols=cols, rows=rows), encoding="utf-8")
        return path
    import plotly.express as px
    import plotly.io as pio

    xs = [r.get(x) for r in rows]
    ys = [r.get(y) for r in rows]
    if kind == "bar":
        fig = px.bar(x=xs, y=ys, color=[r.get(color) for r in rows] if color else None, labels={"x": x, "y": y}, title=title)
    elif kind == "line":
        fig = px.line(x=xs, y=ys, labels={"x": x, "y": y}, title=title, markers=True)
    elif kind == "scatter":
        fig = px.scatter(x=xs, y=ys, size=[r.get(size) for r in rows] if size else None, color=[r.get(color) for r in rows] if color else None, labels={"x": x, "y": y}, title=title)
    elif kind == "heatmap":
        xs_u = sorted({str(v) for v in xs})
        ys_u = sorted({str(v) for v in ys})
        cell = {(str(r.get(y)), str(r.get(x))): r.get(color) for r in rows}
        fig = px.imshow([[cell.get((yy, xx)) for xx in xs_u] for yy in ys_u], x=xs_u, y=ys_u, labels={"x": x, "y": y, "color": color}, title=title, aspect="auto")
    else:
        raise ValueError(kind)
    fig.update_layout(xaxis_tickangle=-45)
    pio.write_html(fig, str(path), include_plotlyjs="cdn", full_html=True, div_id="chart")
    return path


# ---- the executor ---------------------------------------------------------------------------


def run_chain(ctx: Ctx, chain: list[list[Any]], out_dir: Path, name: str) -> tuple[list[Path], list[str], Rows]:
    rows: Rows = []
    saved: dict[str, Rows] = {}
    files: list[Path] = []
    shows: list[str] = []
    for step in chain:
        op, params = (step[0], step[1] if len(step) > 1 else {})
        params = dict(params)
        if op in ("facts", "required_actions", "ordered_steps", "forced_edges", "anchors", "definitions", "prohibitions", "who_must_what", "records", "cases", "events", "minutes_sentences", "minutes_sections", "participants", "attendance", "decisions", "tags", "conformance", "variants", "dfg", "handover", "execution_trace", "roundtrips", "bottlenecks"):
            rows = ctx.source(op, params.get("of", ""))
        elif op == "filter":
            rows = op_filter(rows, params["where"])
        elif op == "select":
            rows = op_select(rows, params["fields"])
        elif op == "derive":
            rows = op_derive(rows, params.pop("field"), params.pop("fn"), **params)
        elif op == "group":
            rows = op_group(rows, params.get("by", []), params["aggregates"])
        elif op == "join":
            rows = op_join(rows, saved[params["with"]], params.get("keys") or params.get("on"), params.get("how", "inner"))
        elif op == "sort":
            rows = op_sort(rows, params["by"], params.get("desc", False))
        elif op == "limit":
            rows = op_limit(rows, params["n"])
        elif op == "distinct":
            rows = op_distinct(rows, params.get("fields"))
        elif op == "pivot":
            rows = op_pivot(rows, params["row"], params["col"], params.get("value", "*"), params.get("fn", "count"))
        elif op == "save_as":
            saved[params["name"]] = list(rows)
        elif op == "clingo":
            rows = op_clingo(rows, params["program"], params["in"], params["fields"], params["out"], params["out_fields"])
        elif op == "decision_table":
            rows, p = op_decision_table(rows, params["name"], params["inputs"], params["outputs"], params["rules"], out_dir)
            files.append(p)
        elif op == "workbook":
            files.append(render_workbook(rows, out_dir / f"{params.get('name', slug(name))}.xlsx", params.get("sheet", "rows")))
        elif op == "document":
            files.append(render_document(rows, out_dir / f"{params.get('name', slug(name))}.docx", params.get("title", name), params.get("columns")))
        elif op == "checklist":
            files.append(render_checklist(rows, out_dir / f"{params.get('name', slug(name))}.docx", params.get("title", name), params["item"], params.get("note")))
        elif op == "page":
            files.append(render_page(rows, out_dir / f"{params.get('name', slug(name))}.html", params.get("title", name), params.get("kind", "table"), params.get("x", ""), params.get("y", ""), params.get("color", ""), params.get("size", "")))
        elif op == "json":
            p = out_dir / f"{params.get('name', slug(name))}.json"
            p.write_text(json.dumps(rows, indent=1, sort_keys=True, ensure_ascii=False, default=str) + "\n", encoding="utf-8")
            files.append(p)
        elif op == "bpmn":
            files.append(ctx.h["bpmn_from_edges"](rows, out_dir / f"{params.get('name', slug(name))}.bpmn"))
        else:
            raise ValueError(op)
    shows.append(f"rows {len(rows)}" + (f"; first {json.dumps(rows[0], default=str, ensure_ascii=False)[:220]}" if rows else ""))
    return files, shows, rows


def arrows(chain: list[list[Any]]) -> list[str]:
    out = []
    for step in chain:
        op, params = (step[0], step[1] if len(step) > 1 else {})
        detail = ""
        if params:
            detail = " " + ", ".join(f"{k}={json.dumps(v, ensure_ascii=False) if not isinstance(v, str) else v}" for k, v in params.items() if k not in ("program", "rules"))
            if "program" in params:
                detail += " program: " + " ".join(params["program"].split())[:160]
        out.append(f"-> {op}{detail} ({BINDINGS.get(op, '?')})")
    return out


def run_register(ctx: Ctx, register: Path) -> dict[str, Any]:
    reg = yaml.safe_load(register.read_text(encoding="utf-8"))
    report: dict[str, Any] = {"practices": []}
    for practice in reg["practices"]:
        pr = {"name": practice["name"], "group": practice.get("group", ""), "deliverables": []}
        for d in practice["deliverables"]:
            entry = {"name": d["name"], "shape": d.get("shape", ""), "instantiated_on": d.get("instantiated_on", "none"), "chain": d.get("chain", []), "files": [], "shows": [], "status": ""}
            if d.get("instantiated_on", "none") == "none" or not d.get("chain"):
                entry["status"] = "no real input"
            else:
                out_dir = OUT / slug(practice["name"]) / slug(d["name"])
                out_dir.mkdir(parents=True, exist_ok=True)
                try:
                    files, shows, rows = run_chain(ctx, d["chain"], out_dir, d["name"])
                    for f in files:
                        canonical_office(f)
                    entry["files"] = [(str(p.relative_to(ROOT)), sha256_file(p), check_register(str(p.relative_to(ROOT)), sha256_file(p))) for p in files]
                    entry["shows"] = shows
                    entry["status"] = "proven" if files and rows else ("empty" if files else "failed")
                except Exception as e:  # the failure is the evidence
                    entry["status"] = "failed"
                    entry["shows"] = [f"{type(e).__name__}: {str(e)[:200]}"]
            pr["deliverables"].append(entry)
        report["practices"].append(pr)
    return report


def md_section(report: dict[str, Any]) -> list[str]:
    lines = ["## Baseline: the ITIL surface", "", "```"]
    files = [f for pr in report["practices"] for d in pr["deliverables"] for f in d["files"]]
    st = collections.Counter(f[2] for f in files)
    lines.append(f"deliverable files {len(files)}: reproduced {st['reproduced']}, registered now {st['registered']}, changed {st['CHANGED']}")
    lines.append("")
    lines.append(f"{'practice':45s} {'deliverables':>12s} {'proven':>7s} {'empty':>6s} {'failed':>7s} {'no real input':>14s}")
    tot = collections.Counter()
    for pr in report["practices"]:
        c = collections.Counter(d["status"] for d in pr["deliverables"])
        tot.update(c)
        lines.append(f"{pr['name'][:45]:45s} {len(pr['deliverables']):12d} {c['proven']:7d} {c['empty']:6d} {c['failed']:7d} {c['no real input']:14d}")
    lines.append(f"{'all':45s} {sum(tot.values()):12d} {tot['proven']:7d} {tot['empty']:6d} {tot['failed']:7d} {tot['no real input']:14d}")
    lines += ["```", ""]
    for pr in report["practices"]:
        lines += [f"### {pr['name']}", ""]
        for d in pr["deliverables"]:
            lines += [f"**{d['name']}**  {d['status']}", "", "```"]
            if d["shape"]:
                lines.append(f"shape: {d['shape']}")
            lines.append(f"instantiated on: {d['instantiated_on']}")
            lines += arrows(d["chain"])
            lines.append(f"-> {d['name']}")
            for f, h, st in d["files"]:
                lines.append(f"{f:70s} sha256 {h}  {st}")
            for s in d["shows"]:
                lines.append(f"shows: {s}")
            lines += ["```", ""]
    return lines
